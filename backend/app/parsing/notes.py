"""Time-window notes parser + validation (deterministic; no LLM).

Handles well-formed CSV / XLSX / DOCX notes and finds the columns it needs
(Nr, Start, Stop, GPS, light/dark, location) **intelligently** — by exact header
token *and* by substring keyword, in English and simple Polish — so verbose
real-world headers resolve without a hard-coded alias (``Start``/``End``,
``Light or Dark``, ``Other site Info``, …); unrelated columns (``Date``,
``Type Of Measurement``, ``TEMPERATURA``) are ignored. The delimiter is
auto-detected (tab / ``;`` / ``,`` / runs of 2+ spaces for aligned exports). It
normalises clean time formats (``9.38``, ``9:38``, ``09:38:00`` → ``HH:MM:SS``)
but does **not** try to repair genuinely messy handwriting — that is the smart
feature's job. Unparseable values are left blank and surfaced by ``validate_notes``.

# TODO: LLM tolerant parsing of messy field notes (seminar 6). This
# deterministic parser covers well-formed files.
"""

import re
from dataclasses import dataclass, field
from pathlib import Path

import pandas as pd

from app.parsing.encoding import detect_encoding

# Flag values mirror the frontend's NoteFlag union.
STOP_BEFORE_START = "stop_before_start"
GPS_MISSING = "gps_missing"
UNPARSEABLE_TIME = "unparseable_time"
LOCATION_MISSING = "location_missing"

# Canonical column -> (exact whole-header tokens, substring keywords), both
# matched against the header with all whitespace stripped and lower-cased.
# *Exact* tokens are for short/ambiguous words that must match the whole header
# (``od``/``do``/``type``/``end``), so they can't fire inside a longer name.
# *Substring* keywords match anywhere, so a verbose header like ``Light or Dark``
# or ``Other site Info`` still resolves. Order = resolution priority; each header
# is claimed by the first field it matches, so specific fields win over the broad
# location keywords. ``type`` maps to light/dark (a common column name) but only
# as an exact match, so ``Type Of Measurement`` is left for no field.
_FIELD_MATCHERS: dict[str, tuple[frozenset[str], tuple[str, ...]]] = {
    "nr": (
        frozenset({"nr", "no", "no.", "lp", "lp.", "l.p.", "id", "#"}),
        ("numer", "punkt", "pkt", "spot"),
    ),
    "start": (
        frozenset({"od"}),
        ("start", "begin", "poczate", "począte", "from"),
    ),
    "stop": (
        frozenset({"do", "end"}),
        ("stop", "end", "koniec", "finish"),
    ),
    "light_dark": (
        frozenset({"type", "l/d", "ld"}),
        ("light", "dark", "chamber", "komora", "jasn", "ciemn"),
    ),
    "gps": (
        frozenset({"gps"}),
        ("gps", "coord", "wspol", "współ"),
    ),
    "location": (
        frozenset(),
        (
            "location",
            "site",
            "info",
            "opis",
            "miejsce",
            "lokal",
            "gdzie",
            "descr",
            "comment",
            "koment",
            "uwag",
            "note",
            "other",
        ),
    ),
}


@dataclass
class NoteRow:
    """One parsed notes row. ``flags`` is filled by ``validate_notes``."""

    nr: int
    start_time: str  # "HH:MM:SS", or "" if unparseable
    stop_time: str
    gps: str
    light_dark: str  # "light" | "dark" | ""
    location: str
    flags: list[str] = field(default_factory=list)


def normalize_time(raw: str) -> str:
    """Normalise a clean time to ``HH:MM:SS``; return ``""`` if unparseable."""
    s = raw.strip().replace(".", ":")
    if not s:
        return ""
    parts = s.split(":")
    try:
        nums = [int(p) for p in parts]
    except ValueError:
        return ""
    if len(nums) == 2:
        hour, minute, second = nums[0], nums[1], 0
    elif len(nums) == 3:
        hour, minute, second = nums
    else:
        return ""
    if not (0 <= hour < 24 and 0 <= minute < 60 and 0 <= second < 60):
        return ""
    return f"{hour:02d}:{minute:02d}:{second:02d}"


def _normalize_light_dark(raw: str) -> str:
    s = raw.strip().lower()
    if not s:
        return ""
    if s[0] in {"l", "j"}:  # light / jasny
        return "light"
    if s[0] in {"d", "c"}:  # dark / ciemny
        return "dark"
    return ""


def _parse_nr(raw: str) -> int:
    """Leading integer of the Nr cell (tolerates a trailing redo marker like ``3!``)."""
    match = re.match(r"\s*(\d+)", raw)
    return int(match.group(1)) if match else 0


def _read_table(path: str | Path) -> list[dict[str, str]]:
    """Read the notes file into a list of string-keyed/valued row dicts."""
    suffix = Path(path).suffix.lower()
    if suffix in {".csv", ".txt", ".tsv"}:
        frame = _read_csv_autodetect(path)
    elif suffix in {".xlsx", ".xls"}:
        frame = pd.read_excel(path, dtype=str, engine="openpyxl").fillna("")
    elif suffix == ".docx":
        return _read_docx_table(path)
    else:
        raise ValueError(f"Unsupported notes format: {suffix!r}")
    return [
        {str(k): "" if v is None else str(v) for k, v in record.items()}
        for record in frame.to_dict("records")
    ]


def _read_csv_autodetect(path: str | Path) -> pd.DataFrame:
    """Read a delimited notes file, auto-detecting encoding then delimiter.

    Encoding is sniffed first (Windows exports are often cp1250/UTF-16, and a
    site name like ``nad tamą`` carries the non-UTF-8 byte that used to reject the
    whole file); delimiter is tab / ``;`` / ``,`` / runs of 2+ spaces.
    """
    encoding = detect_encoding(path)
    for sep in ("\t", ";", ","):
        try:
            frame = pd.read_csv(
                path, dtype=str, keep_default_na=False, sep=sep, encoding=encoding
            )
        except ValueError:  # pandas ParserError subclasses ValueError
            continue
        if frame.shape[1] > 1:
            return frame
    # Space-aligned (fixed-width-ish) notes: split on runs of 2+ spaces so verbose
    # headers ("Type Of Measurement") and values ("nad tamą") keep their single
    # internal spaces instead of being torn into separate columns.
    try:
        frame = pd.read_csv(
            path,
            dtype=str,
            keep_default_na=False,
            sep=r"\s{2,}",
            engine="python",
            encoding=encoding,
        )
        if frame.shape[1] > 1:
            return frame
    except ValueError:
        pass
    return pd.read_csv(
        path,
        dtype=str,
        keep_default_na=False,
        sep=None,
        engine="python",
        encoding=encoding,
    )


def _read_docx_table(path: str | Path) -> list[dict[str, str]]:
    """Read the first table of a Word document into row dicts (header = row 0)."""
    from docx import Document

    document = Document(str(path))
    if not document.tables:
        return []
    table = document.tables[0]
    header = [cell.text.strip() for cell in table.rows[0].cells]
    rows: list[dict[str, str]] = []
    for row in table.rows[1:]:
        values = [cell.text.strip() for cell in row.cells]
        rows.append(dict(zip(header, values, strict=False)))
    return rows


def _resolve_columns(headers: list[str]) -> dict[str, str]:
    """Map each canonical field to the actual header present in the file.

    Matching is by exact whole-header token *and* by substring keyword (see
    ``_FIELD_MATCHERS``), so verbose real-world headers resolve without an exact
    alias: ``Start``/``End`` → start/stop, ``Light or Dark`` → light/dark,
    ``Other site Info`` → location, ``GPS`` → gps, while ``Type Of Measurement``
    (only an exact ``type`` counts) and ``Date``/``TEMPERATURA`` stay unmapped.
    All whitespace/newlines are stripped first (a Word table can wrap a cell as
    ``"Light\\n/dark"``). Fields resolve in priority order and each header is
    claimed once, so specific fields win over the broad location keywords.
    """
    resolved: dict[str, str] = {}
    claimed: set[str] = set()
    for canonical, (exact, keywords) in _FIELD_MATCHERS.items():
        for header in headers:
            if header in claimed:
                continue
            key = re.sub(r"\s+", "", header).lower()
            if key in exact or any(kw in key for kw in keywords):
                resolved[canonical] = header
                claimed.add(header)
                break
    return resolved


def parse_notes(path: str | Path) -> list[NoteRow]:
    """Parse a notes file into ``NoteRow``s (flags empty — call ``validate_notes``)."""
    records = _read_table(path)
    if not records:
        return []
    columns = _resolve_columns(list(records[0].keys()))

    def cell(record: dict[str, str], canonical: str) -> str:
        source = columns.get(canonical)
        return record.get(source, "").strip() if source else ""

    rows: list[NoteRow] = []
    for record in records:
        rows.append(
            NoteRow(
                nr=_parse_nr(cell(record, "nr")),
                start_time=normalize_time(cell(record, "start")),
                stop_time=normalize_time(cell(record, "stop")),
                gps=cell(record, "gps"),
                light_dark=_normalize_light_dark(cell(record, "light_dark")),
                location=cell(record, "location"),
            )
        )

    # ``nr`` is the app's unique per-spot key. If the file had no recognisable
    # number column (all 0) or the numbers collide (e.g. a "4" and a "4.5" dark
    # pair both truncate to 4), fall back to a 1-based row index so every spot
    # stays distinct. light/dark is preserved separately.
    nrs = [r.nr for r in rows]
    if any(n == 0 for n in nrs) or len(set(nrs)) != len(nrs):
        for index, row in enumerate(rows, start=1):
            row.nr = index
    return rows


def validate_notes(rows: list[NoteRow]) -> list[NoteRow]:
    """Set each row's ``flags`` in place (and return the list) for the Confirm step."""
    for row in rows:
        flags: list[str] = []
        if not row.start_time or not row.stop_time:
            flags.append(UNPARSEABLE_TIME)
        elif row.stop_time <= row.start_time:
            # "HH:MM:SS" strings compare chronologically within a day.
            flags.append(STOP_BEFORE_START)
        # Blank or a bare "?" (the field shorthand for "no fix") counts as missing.
        if not row.gps.strip() or row.gps.strip() == "?":
            flags.append(GPS_MISSING)
        if not row.location.strip():
            flags.append(LOCATION_MISSING)
        row.flags = flags
    return rows
